#!/usr/bin/env python3
"""Interview transcript -> summary fusion pipeline.

Subcommands:
1) prepare  : convert docx to temp markdown + extract anchors/chunks/json input
2) run-llm  : call OpenAI-compatible chat completion API in batches
3) apply    : write model results back into docx (reply paragraphs), mark red
4) pipeline : run prepare -> run-llm -> apply in one command
5) agent-pack          : generate Codex/Claude-ready batch tasks
6) merge-agent-results : merge batch outputs into unified results json
7) prepare-agent       : run prepare + agent-pack
8) finalize-agent      : run merge-agent-results + apply
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.shared import RGBColor


RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
LOW_CONF_NOTE = "（注：低置信度，需要人工审核）"

SPEAKER_RE = re.compile(
    r"^(?P<speaker>.+?)\((?P<timestamp>\d{2}:\d{2}:\d{2})\)\s*[：:]\s*(?P<text>.*)$"
)
SECTION_HEADING_RE = re.compile(r"^[一二三四五六七八九十]+、")
SUBSECTION_HEADING_RE = re.compile(r"^\d+、[^？?]{1,40}$")

SYSTEM_PROMPT = """
你是“访谈纪要融合引擎”。你的任务是将转写内容中的回答，按语义匹配填入访谈纪要问题对应的“回复”。

强制规则：
1. 语义匹配优先，关键词仅辅助。
2. 回答必须可追溯到证据；无证据不填。
3. 同一问题多段分散回答可合并为 2-4 句精炼回复。
4. 无明确对应回答时，留空（status=blank）。
5. 不编造，不补充外部事实。
6. 支持“一次提问覆盖多个纪要问题”的情况：先拆分回答要点，再按纪要问题意图分别落位，避免机械复制同一答案。
7. 对于同一 question_text 的多个回复槽位（slot_index_for_question / slot_total_for_question），应按槽位语义做差异化分配；若无额外信息可留空，不强行重复填充。
8. 若转写中出现纪要里没有的“临时新增问题”，可输出到 new_items（含 question_text、answer_draft、evidence、confidence）。
9. 冲突裁决顺序：
   a) 直接回答核心意图优先；
   b) 信息完整、可执行性高优先；
   c) 互补可合并；
   d) 后续明确修正优先；
   e) 更具体事实细节优先；
   f) 仍无法判定则留空。

低置信度判定：
- 硬规则（命中即留空）：
  1) 无可追溯证据；
  2) 语义不直接回答问题；
  3) 明显冲突且无法消解。
- 评分：
  confidence = 0.35*S + 0.25*E + 0.20*C + 0.10*F + 0.10*T
  S=语义匹配度, E=证据覆盖度, C=一致性, F=完整性, T=可追溯性（均为0-100）
- 阈值：
  confidence >= 75: filled
  55 <= confidence < 75: filled + note_required=true
  confidence < 55: blank

输出必须是合法 JSON，不要输出解释性文字。
""".strip()

AGENT_WORKER_TEMPLATE = f"""
你正在处理一个访谈纪要融合批次任务。

你必须遵循以下 system prompt 规则：
---
{SYSTEM_PROMPT}
---

请读取当前批次输入 JSON（字段：project_id/questions/transcript_chunks），并仅输出合法 JSON 对象：
{{
  "project_id": "string",
  "results": [
    {{
      "question_id": "string",
      "status": "filled | blank",
      "confidence": 0,
      "evidence": ["T0001"],
      "answer_draft": "string",
      "note_required": false,
      "note_text": ""
    }}
  ],
  "new_items": [
    {{
      "question_text": "string",
      "status": "filled | blank",
      "confidence": 0,
      "evidence": ["T0001"],
      "answer_draft": "string",
      "note_required": false,
      "note_text": ""
    }}
  ]
}}

额外要求：
1. 每个 question_id 必须输出一条结果。
2. answer_draft 必须是 2-4 句中文纪要风格。
3. status=blank 时 answer_draft 必须为空字符串。
4. note_required=true 时 note_text 固定为：{LOW_CONF_NOTE}
5. new_items 只保留纪要中不存在且在转写中有明确证据支撑的问题。
6. 只输出 JSON，不要 markdown，不要解释。
""".strip()


@dataclass
class QuestionAnchor:
    question_id: str
    question_text: str
    question_paragraph_index: int
    reply_paragraph_index: int
    existing_reply: str
    section_title: str
    subsection_title: str
    slot_index_for_question: int
    slot_total_for_question: int

    def to_json(self) -> Dict[str, object]:
        return {
            "question_id": self.question_id,
            "question_text": self.question_text,
            "question_paragraph_index": self.question_paragraph_index,
            "reply_paragraph_index": self.reply_paragraph_index,
            "existing_reply": self.existing_reply,
            "section_title": self.section_title,
            "subsection_title": self.subsection_title,
            "slot_index_for_question": self.slot_index_for_question,
            "slot_total_for_question": self.slot_total_for_question,
        }


def read_nonempty_paragraphs(doc: Document) -> List[Tuple[int, str]]:
    rows: List[Tuple[int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            rows.append((idx, text))
    return rows


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def write_json(path: Path, payload: object) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_transcript_chunks(transcript_doc: Document) -> List[Dict[str, object]]:
    chunks: List[Dict[str, object]] = []
    counter = 1
    for p_idx, p in enumerate(transcript_doc.paragraphs):
        raw = p.text.strip()
        if not raw:
            continue
        m = SPEAKER_RE.match(raw)
        if m:
            speaker = m.group("speaker").strip()
            ts = m.group("timestamp").strip()
            text = m.group("text").strip()
        else:
            speaker = ""
            ts = ""
            text = raw
        chunks.append(
            {
                "chunk_id": f"T{counter:04d}",
                "paragraph_index": p_idx,
                "speaker": speaker,
                "timestamp": ts,
                "text": text,
                "raw": raw,
            }
        )
        counter += 1
    return chunks


def markdown_from_transcript_chunks(chunks: Sequence[Dict[str, object]]) -> str:
    lines = ["# temp_transcript", ""]
    for c in chunks:
        speaker = str(c.get("speaker", "")).strip() or "Unknown"
        ts = str(c.get("timestamp", "")).strip()
        ts_label = f"[{ts}] " if ts else ""
        text = str(c.get("text", "")).strip()
        lines.append(
            f"- {c['chunk_id']} {ts_label}{speaker}: {text}".rstrip()
        )
    lines.append("")
    return "\n".join(lines)


def markdown_from_summary_doc(summary_doc: Document) -> str:
    lines = ["# temp_summary", "", "## Paragraphs", ""]
    for idx, text in read_nonempty_paragraphs(summary_doc):
        lines.append(f"- P{idx:04d}: {text}")

    for t_idx, table in enumerate(summary_doc.tables):
        lines.extend(["", f"## Table {t_idx}", ""])
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            lines.append(f"- R{r_idx:03d}: " + " | ".join(cells))
    lines.append("")
    return "\n".join(lines)


def is_reply_line(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("回复：") or stripped.startswith("回复:")


def split_reply_prefix_tail(text: str) -> Tuple[str, str]:
    stripped = text.strip()
    m = re.match(r"^回复[：:]\s*", stripped)
    if m:
        prefix = m.group(0)
        return prefix, stripped[m.end() :].strip()
    if stripped == "回复":
        return "回复：", ""
    if stripped.startswith("回复"):
        rest = stripped[len("回复") :].lstrip("：: ")
        return "回复：", rest
    return "回复：", stripped


def is_heading_line(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if t in {"请问：", "请问:", "资料清单"}:
        return True
    if re.match(r"^声\s*明$", t):
        return True
    if SECTION_HEADING_RE.match(t):
        return True
    if SUBSECTION_HEADING_RE.match(t):
        return True
    return False


def find_context_titles(paragraph_texts: Sequence[str], idx: int) -> Tuple[str, str]:
    section_title = ""
    subsection_title = ""
    for i in range(idx, -1, -1):
        text = paragraph_texts[i].strip()
        if not text:
            continue
        if not section_title and SECTION_HEADING_RE.match(text):
            section_title = text
        if not subsection_title and SUBSECTION_HEADING_RE.match(text):
            subsection_title = text
        if section_title and subsection_title:
            break
    return section_title, subsection_title


def find_question_before_reply(paragraph_texts: Sequence[str], reply_idx: int) -> Tuple[int, str]:
    fallback: Optional[Tuple[int, str]] = None
    for i in range(reply_idx - 1, -1, -1):
        text = paragraph_texts[i].strip()
        if not text:
            continue
        if is_reply_line(text):
            continue
        if is_heading_line(text):
            if fallback is not None:
                break
            continue
        if "？" in text or "?" in text:
            return i, text
        if fallback is None:
            fallback = (i, text)
    return fallback if fallback is not None else (-1, "")


def extract_question_anchors(summary_doc: Document) -> List[QuestionAnchor]:
    p_texts = [p.text.strip() for p in summary_doc.paragraphs]
    anchors_raw: List[QuestionAnchor] = []
    q_counter = 1
    for idx, text in enumerate(p_texts):
        if not is_reply_line(text):
            continue
        q_idx, q_text = find_question_before_reply(p_texts, idx)
        _, existing = split_reply_prefix_tail(text)
        section_title, subsection_title = find_context_titles(p_texts, idx)
        anchors_raw.append(
            QuestionAnchor(
                question_id=f"Q{q_counter:04d}",
                question_text=q_text,
                question_paragraph_index=q_idx,
                reply_paragraph_index=idx,
                existing_reply=existing,
                section_title=section_title,
                subsection_title=subsection_title,
                slot_index_for_question=1,
                slot_total_for_question=1,
            )
        )
        q_counter += 1

    group_counts: Dict[Tuple[str, str, str], int] = {}
    for a in anchors_raw:
        key = (a.question_text, a.section_title, a.subsection_title)
        group_counts[key] = group_counts.get(key, 0) + 1

    running_index: Dict[Tuple[str, str, str], int] = {}
    anchors: List[QuestionAnchor] = []
    for a in anchors_raw:
        key = (a.question_text, a.section_title, a.subsection_title)
        running_index[key] = running_index.get(key, 0) + 1
        a.slot_index_for_question = running_index[key]
        a.slot_total_for_question = group_counts[key]
        anchors.append(a)
    return anchors


def build_llm_input(
    project_id: str,
    anchors: Sequence[QuestionAnchor],
    chunks: Sequence[Dict[str, object]],
) -> Dict[str, object]:
    questions = []
    for a in anchors:
        questions.append(
            {
                "question_id": a.question_id,
                "question_text": a.question_text,
                "existing_reply": a.existing_reply,
                "question_paragraph_index": a.question_paragraph_index,
                "reply_paragraph_index": a.reply_paragraph_index,
                "section_title": a.section_title,
                "subsection_title": a.subsection_title,
                "slot_index_for_question": a.slot_index_for_question,
                "slot_total_for_question": a.slot_total_for_question,
            }
        )
    transcript_chunks = []
    for c in chunks:
        transcript_chunks.append(
            {
                "chunk_id": c["chunk_id"],
                "speaker": c.get("speaker", ""),
                "text": c.get("text", ""),
                "position": f"P{int(c.get('paragraph_index', 0)):04d}",
            }
        )
    return {
        "project_id": project_id,
        "questions": questions,
        "transcript_chunks": transcript_chunks,
    }


def prepare(args: argparse.Namespace) -> None:
    output_dir = Path(args.output_dir)
    ensure_dir(output_dir)

    transcript_doc = Document(str(args.transcript_docx))
    summary_doc = Document(str(args.summary_docx))

    chunks = parse_transcript_chunks(transcript_doc)
    anchors = extract_question_anchors(summary_doc)
    project_id = args.project_id or Path(args.summary_docx).stem

    (output_dir / "temp_transcript.md").write_text(
        markdown_from_transcript_chunks(chunks),
        encoding="utf-8",
    )
    (output_dir / "temp_summary.md").write_text(
        markdown_from_summary_doc(summary_doc),
        encoding="utf-8",
    )
    write_json(output_dir / "transcript_chunks.json", chunks)
    write_json(output_dir / "questions.json", [a.to_json() for a in anchors])
    write_json(output_dir / "llm_input.json", build_llm_input(project_id, anchors, chunks))

    print(f"[prepare] project_id={project_id}")
    print(f"[prepare] transcript chunks: {len(chunks)}")
    print(f"[prepare] question anchors: {len(anchors)}")
    print(f"[prepare] outputs in: {output_dir}")


def safe_json_load_text(text: str) -> Dict[str, object]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```[a-zA-Z0-9]*\n?", "", cleaned)
        cleaned = re.sub(r"\n?```$", "", cleaned.strip())
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, dict):
            return obj
    except json.JSONDecodeError:
        pass

    m = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
    if not m:
        raise ValueError("LLM output does not contain a JSON object.")
    obj = json.loads(m.group(0))
    if not isinstance(obj, dict):
        raise ValueError("Parsed JSON root is not an object.")
    return obj


def chat_completions_url(base_url: str) -> str:
    base = base_url.rstrip("/")
    if base.endswith("/chat/completions"):
        return base
    if not base.endswith("/v1"):
        base += "/v1"
    return base + "/chat/completions"


def call_llm_once(
    *,
    api_key: str,
    base_url: str,
    model: str,
    system_prompt: str,
    user_prompt: str,
    temperature: float,
    timeout: int,
) -> str:
    url = chat_completions_url(base_url)
    payload = {
        "model": model,
        "temperature": temperature,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        raw = resp.read().decode("utf-8")
    body = json.loads(raw)
    content = body["choices"][0]["message"]["content"]
    if not isinstance(content, str):
        raise ValueError("Unexpected LLM response content type.")
    return content


def build_user_prompt(batch_payload: Dict[str, object]) -> str:
    return (
        "请处理以下输入数据，输出 JSON 对象，键名为 project_id、results、new_items。\n\n"
        "要求：\n"
        "1. 每个 question_id 都要输出一条。\n"
        "2. 支持一次提问覆盖多个纪要问题，先拆要点再分配到各 question_id。\n"
        "3. answer_draft 为 2-4 句中文，纪要风格。\n"
        "4. status=blank 时 answer_draft 为空字符串。\n"
        "5. evidence 给出命中的 chunk_id 列表；blank 可为空数组。\n"
        f"6. note_required=true 时 note_text 固定写：{LOW_CONF_NOTE}\n"
        "7. 可输出 new_items（纪要外临时问题），仅限有明确证据支持的项。\n"
        "8. 只输出合法 JSON，不要 markdown。\n\n"
        "输入数据：\n"
        + json.dumps(batch_payload, ensure_ascii=False, indent=2)
    )


def build_batch_payload(
    project_id: str,
    questions_batch: Sequence[Dict[str, object]],
    chunks: Sequence[Dict[str, object]],
) -> Dict[str, object]:
    return {
        "project_id": project_id,
        "questions": list(questions_batch),
        "transcript_chunks": list(chunks),
    }


def normalize_one_result(
    question_id: str, row: Dict[str, object]
) -> Dict[str, object]:
    status_raw = str(row.get("status", "blank")).strip().lower()
    status = "filled" if status_raw == "filled" else "blank"
    try:
        confidence = float(row.get("confidence", 0))
    except (TypeError, ValueError):
        confidence = 0.0
    confidence = max(0.0, min(100.0, confidence))

    answer = str(row.get("answer_draft", "") or "").strip()
    evidence_raw = row.get("evidence", [])
    evidence: List[str] = []
    if isinstance(evidence_raw, list):
        for item in evidence_raw:
            if isinstance(item, (str, int, float)):
                evidence.append(str(item))

    note_required = bool(row.get("note_required", False))
    note_text = str(row.get("note_text", "") or "").strip()

    # Enforce PRD threshold policy.
    if confidence < 55:
        status = "blank"
    elif confidence < 75 and status == "filled":
        note_required = True

    if status == "blank":
        answer = ""
        note_required = False
        note_text = ""
    elif note_required and not note_text:
        note_text = LOW_CONF_NOTE

    return {
        "question_id": question_id,
        "status": status,
        "confidence": round(confidence, 2),
        "evidence": evidence,
        "answer_draft": answer,
        "note_required": note_required,
        "note_text": note_text,
    }


def normalize_text_key(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip()


def normalize_new_item(row: Dict[str, object]) -> Optional[Dict[str, object]]:
    question_text = str(row.get("question_text", "") or "").strip()
    if not question_text:
        return None

    status_raw = str(row.get("status", "filled")).strip().lower()
    status = "filled" if status_raw == "filled" else "blank"
    try:
        confidence = float(row.get("confidence", 0))
    except (TypeError, ValueError):
        confidence = 0.0
    confidence = max(0.0, min(100.0, confidence))

    answer = str(row.get("answer_draft", "") or "").strip()
    evidence_raw = row.get("evidence", [])
    evidence: List[str] = []
    if isinstance(evidence_raw, list):
        for item in evidence_raw:
            if isinstance(item, (str, int, float)):
                evidence.append(str(item))

    note_required = bool(row.get("note_required", False))
    note_text = str(row.get("note_text", "") or "").strip()

    if confidence < 55:
        status = "blank"
    elif confidence < 75 and status == "filled":
        note_required = True

    if status == "blank":
        answer = ""
        note_required = False
        note_text = ""
    elif note_required and not note_text:
        note_text = LOW_CONF_NOTE

    return {
        "question_text": question_text,
        "status": status,
        "confidence": round(confidence, 2),
        "evidence": evidence,
        "answer_draft": answer,
        "note_required": note_required,
        "note_text": note_text,
    }


def chunked(seq: Sequence[Dict[str, object]], size: int) -> Iterable[Sequence[Dict[str, object]]]:
    for i in range(0, len(seq), size):
        yield seq[i : i + size]


def run_llm(args: argparse.Namespace) -> None:
    input_path = Path(args.input_json)
    output_path = Path(args.output_json)
    payload = json.loads(input_path.read_text(encoding="utf-8"))
    project_id = str(payload.get("project_id", "project"))
    questions = payload.get("questions", [])
    chunks = payload.get("transcript_chunks", [])
    if not isinstance(questions, list) or not isinstance(chunks, list):
        raise ValueError("input_json format invalid: questions/transcript_chunks must be arrays.")

    api_key = args.api_key or os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is required (or pass --api-key).")

    base_url = args.base_url or os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
    model = args.model or os.environ.get("OPENAI_MODEL", "gpt-4.1")

    all_results: Dict[str, Dict[str, object]] = {}
    new_items_map: Dict[str, Dict[str, object]] = {}
    total_batches = (len(questions) + args.batch_size - 1) // args.batch_size
    for b_idx, q_batch in enumerate(chunked(questions, args.batch_size), start=1):
        batch_payload = build_batch_payload(project_id, q_batch, chunks)
        user_prompt = build_user_prompt(batch_payload)

        print(f"[run-llm] batch {b_idx}/{total_batches} questions={len(q_batch)}")
        last_error: Optional[Exception] = None
        response_text = ""
        for attempt in range(1, args.max_retries + 1):
            try:
                response_text = call_llm_once(
                    api_key=api_key,
                    base_url=base_url,
                    model=model,
                    system_prompt=SYSTEM_PROMPT,
                    user_prompt=user_prompt,
                    temperature=args.temperature,
                    timeout=args.timeout,
                )
                last_error = None
                break
            except (urllib.error.URLError, urllib.error.HTTPError, ValueError, KeyError, json.JSONDecodeError) as exc:
                last_error = exc
                wait_sec = min(2 * attempt, 8)
                print(f"[run-llm] attempt {attempt} failed: {exc}. retry in {wait_sec}s")
                time.sleep(wait_sec)

        if last_error is not None:
            raise RuntimeError(f"LLM batch failed after retries: {last_error}") from last_error

        parsed = safe_json_load_text(response_text)
        batch_results = parsed.get("results", [])
        if not isinstance(batch_results, list):
            raise ValueError("LLM output missing results array.")
        batch_new_items = parsed.get("new_items", [])
        if batch_new_items is not None and not isinstance(batch_new_items, list):
            raise ValueError("LLM output field new_items must be an array if provided.")

        batch_question_ids = {
            str(item.get("question_id", "")).strip()
            for item in q_batch
            if isinstance(item, dict)
        }
        for row in batch_results:
            if not isinstance(row, dict):
                continue
            qid = str(row.get("question_id", "")).strip()
            if not qid or qid not in batch_question_ids:
                continue
            all_results[qid] = normalize_one_result(qid, row)

        for row in batch_new_items or []:
            if not isinstance(row, dict):
                continue
            item = normalize_new_item(row)
            if item is None:
                continue
            if item["status"] != "filled":
                continue
            key = normalize_text_key(str(item["question_text"]))
            if not key:
                continue
            prev = new_items_map.get(key)
            if prev is None or float(item["confidence"]) > float(prev["confidence"]):
                new_items_map[key] = item

        # Fill missing with blank to keep full coverage.
        for item in q_batch:
            qid = str(item.get("question_id", "")).strip()
            if not qid:
                continue
            if qid not in all_results:
                all_results[qid] = normalize_one_result(
                    qid, {"status": "blank", "confidence": 0}
                )

    ordered_results: List[Dict[str, object]] = []
    for item in questions:
        qid = str(item.get("question_id", "")).strip()
        if not qid:
            continue
        ordered_results.append(
            all_results.get(
                qid, normalize_one_result(qid, {"status": "blank", "confidence": 0})
            )
        )

    ensure_dir(output_path.parent)
    write_json(
        output_path,
        {
            "project_id": project_id,
            "results": ordered_results,
            "new_items": list(new_items_map.values()),
        },
    )
    print(f"[run-llm] wrote: {output_path}")


def to_rel_path(base_dir: Path, file_path: Path) -> str:
    try:
        return str(file_path.relative_to(base_dir))
    except ValueError:
        return str(file_path)


def agent_pack(args: argparse.Namespace) -> None:
    input_path = Path(args.input_json)
    pack_dir = Path(args.pack_dir)
    ensure_dir(pack_dir)
    batches_dir = pack_dir / "batches"
    ensure_dir(batches_dir)

    payload = json.loads(input_path.read_text(encoding="utf-8"))
    project_id = str(payload.get("project_id", "project"))
    questions = payload.get("questions", [])
    chunks = payload.get("transcript_chunks", [])
    if not isinstance(questions, list) or not isinstance(chunks, list):
        raise ValueError("input_json format invalid: questions/transcript_chunks must be arrays.")

    manifest_batches: List[Dict[str, object]] = []
    total_batches = (len(questions) + args.batch_size - 1) // args.batch_size
    for b_idx, q_batch in enumerate(chunked(questions, args.batch_size), start=1):
        batch_id = f"B{b_idx:03d}"
        batch_input_path = batches_dir / f"{batch_id}.input.json"
        batch_output_path = batches_dir / f"{batch_id}.output.json"
        batch_payload = build_batch_payload(project_id, q_batch, chunks)
        write_json(batch_input_path, batch_payload)

        question_ids: List[str] = []
        for q in q_batch:
            if isinstance(q, dict):
                qid = str(q.get("question_id", "")).strip()
                if qid:
                    question_ids.append(qid)

        manifest_batches.append(
            {
                "batch_id": batch_id,
                "question_count": len(question_ids),
                "question_ids": question_ids,
                "input_file": to_rel_path(pack_dir, batch_input_path),
                "output_file": to_rel_path(pack_dir, batch_output_path),
            }
        )

    manifest = {
        "project_id": project_id,
        "source_input_json": str(input_path),
        "batch_size": args.batch_size,
        "total_questions": len(questions),
        "total_batches": total_batches,
        "batches": manifest_batches,
    }
    manifest_path = pack_dir / "manifest.json"
    write_json(manifest_path, manifest)
    (pack_dir / "system_prompt.txt").write_text(SYSTEM_PROMPT + "\n", encoding="utf-8")
    (pack_dir / "worker_prompt_template.md").write_text(
        AGENT_WORKER_TEMPLATE + "\n", encoding="utf-8"
    )

    task_lines = [
        "# Agent Tasks",
        "",
        "你正在使用 Codex/Claude Code 批量处理访谈纪要融合任务。",
        "",
        "## 输入文件",
        f"- 批次清单: `{to_rel_path(pack_dir, manifest_path)}`",
        "- 系统提示词: `system_prompt.txt`",
        "- 工作提示词模板: `worker_prompt_template.md`",
        "",
        "## 执行要求",
        "1. 按 `manifest.json` 的 batches 顺序处理。",
        "2. 每个批次读取 `input_file`，基于模板生成结果 JSON，写入对应 `output_file`。",
        "3. 输出必须是 JSON 对象，包含键：`project_id`、`results`，可选 `new_items`。",
        "4. 每个 question_id 必须返回一条结果；缺失项会被后续步骤自动按 blank 补齐。",
        "5. 若识别到纪要外临时问题，可写入 `new_items`（需有证据支撑）。",
        "",
        "## 完成后",
        "运行：",
        "```bash",
        f"python3 scripts/interview_fusion.py merge-agent-results --manifest-json {manifest_path}",
        "```",
    ]
    (pack_dir / "AGENT_TASK.md").write_text("\n".join(task_lines) + "\n", encoding="utf-8")

    print(f"[agent-pack] project_id={project_id}")
    print(f"[agent-pack] total_questions={len(questions)}")
    print(f"[agent-pack] total_batches={total_batches}")
    print(f"[agent-pack] pack_dir={pack_dir}")


def merge_agent_results(args: argparse.Namespace) -> None:
    manifest_path = Path(args.manifest_json)
    manifest_dir = manifest_path.parent
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    project_id = str(manifest.get("project_id", "project"))
    batches = manifest.get("batches", [])
    if not isinstance(batches, list):
        raise ValueError("manifest_json invalid: batches must be an array.")

    all_results: Dict[str, Dict[str, object]] = {}
    new_items_map: Dict[str, Dict[str, object]] = {}
    ordered_qids: List[str] = []
    missing_outputs: List[str] = []

    for batch in batches:
        if not isinstance(batch, dict):
            continue
        output_file = str(batch.get("output_file", "")).strip()
        if not output_file:
            continue
        qids = batch.get("question_ids", [])
        batch_qids: List[str] = []
        if isinstance(qids, list):
            for qid in qids:
                q = str(qid).strip()
                if q:
                    batch_qids.append(q)
                    ordered_qids.append(q)

        output_path = manifest_dir / output_file
        if not output_path.exists():
            missing_outputs.append(str(output_path))
            if args.strict:
                raise FileNotFoundError(f"Missing batch output: {output_path}")
            for qid in batch_qids:
                all_results[qid] = normalize_one_result(
                    qid, {"status": "blank", "confidence": 0}
                )
            continue

        parsed = safe_json_load_text(output_path.read_text(encoding="utf-8"))
        batch_results = parsed.get("results", [])
        if not isinstance(batch_results, list):
            raise ValueError(f"Invalid batch output (missing results array): {output_path}")
        batch_new_items = parsed.get("new_items", [])
        if batch_new_items is not None and not isinstance(batch_new_items, list):
            raise ValueError(f"Invalid batch output (new_items must be an array): {output_path}")

        allowed_qids = set(batch_qids)
        for row in batch_results:
            if not isinstance(row, dict):
                continue
            qid = str(row.get("question_id", "")).strip()
            if not qid or qid not in allowed_qids:
                continue
            all_results[qid] = normalize_one_result(qid, row)

        for qid in batch_qids:
            if qid not in all_results:
                all_results[qid] = normalize_one_result(
                    qid, {"status": "blank", "confidence": 0}
                )

        for row in batch_new_items or []:
            if not isinstance(row, dict):
                continue
            item = normalize_new_item(row)
            if item is None:
                continue
            if item["status"] != "filled":
                continue
            key = normalize_text_key(str(item["question_text"]))
            if not key:
                continue
            prev = new_items_map.get(key)
            if prev is None or float(item["confidence"]) > float(prev["confidence"]):
                new_items_map[key] = item

    # de-dup while preserving order
    seen: set = set()
    uniq_ordered_qids: List[str] = []
    for qid in ordered_qids:
        if qid not in seen:
            seen.add(qid)
            uniq_ordered_qids.append(qid)

    ordered_results = [all_results[qid] for qid in uniq_ordered_qids if qid in all_results]
    new_items = list(new_items_map.values())
    output_json = Path(args.output_json)
    ensure_dir(output_json.parent)
    write_json(
        output_json,
        {"project_id": project_id, "results": ordered_results, "new_items": new_items},
    )

    filled = sum(1 for r in ordered_results if r.get("status") == "filled")
    blank = sum(1 for r in ordered_results if r.get("status") == "blank")
    flagged = sum(1 for r in ordered_results if bool(r.get("note_required", False)))

    print(f"[merge-agent-results] wrote: {output_json}")
    print(
        f"[merge-agent-results] total={len(ordered_results)} filled={filled} "
        f"blank={blank} flagged={flagged} new_items={len(new_items)}"
    )
    if missing_outputs:
        print(f"[merge-agent-results] missing outputs treated as blank: {len(missing_outputs)}")
        for p in missing_outputs[:10]:
            print(f"  - {p}")


def merge_segments(segments: Sequence[Tuple[str, bool]]) -> List[Tuple[str, bool]]:
    merged: List[Tuple[str, bool]] = []
    for text, is_red in segments:
        if not text:
            continue
        if merged and merged[-1][1] == is_red:
            merged[-1] = (merged[-1][0] + text, is_red)
        else:
            merged.append((text, is_red))
    return merged


def diff_segments(old: str, new: str) -> List[Tuple[str, bool]]:
    if not old.strip():
        return [(new, True)] if new else []
    if old == new:
        return [(new, False)] if new else []
    sm = SequenceMatcher(None, old, new)
    out: List[Tuple[str, bool]] = []
    for tag, _, _, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            out.append((new[j1:j2], False))
        elif tag in {"replace", "insert"}:
            out.append((new[j1:j2], True))
        elif tag == "delete":
            continue
    return merge_segments(out)


def copy_run_style(src_run, dst_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.style is not None:
        dst_run.style = src_run.style
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.bold is not None:
        dst_run.font.bold = src_run.font.bold
    if src_run.font.italic is not None:
        dst_run.font.italic = src_run.font.italic
    if src_run.font.underline is not None:
        dst_run.font.underline = src_run.font.underline


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def rewrite_reply_paragraph(paragraph, prefix: str, old_tail: str, new_tail: str) -> None:
    template_run = paragraph.runs[0] if paragraph.runs else None
    segments: List[Tuple[str, bool]] = [(prefix, False)]
    if new_tail:
        if old_tail:
            segments.extend(diff_segments(old_tail, new_tail))
        else:
            segments.append((new_tail, True))
    segments = merge_segments(segments)

    clear_paragraph_runs(paragraph)
    for text, is_red in segments:
        run = paragraph.add_run(text)
        if template_run is not None:
            copy_run_style(template_run, run)
        run.font.color.rgb = RED if is_red else BLACK


def build_new_reply_text(result: Dict[str, object]) -> str:
    if str(result.get("status", "blank")) != "filled":
        return ""
    answer = str(result.get("answer_draft", "") or "").strip()
    if not answer:
        return ""
    note_required = bool(result.get("note_required", False))
    note_text = str(result.get("note_text", "") or "").strip()
    if note_required:
        if not note_text:
            note_text = LOW_CONF_NOTE
        if note_text not in answer:
            answer = f"{answer}{note_text}"
    return answer


def find_insert_anchor_paragraph(doc: Document):
    for p in doc.paragraphs:
        if p.text.strip().startswith("资料清单"):
            return p
    return doc.paragraphs[-1] if doc.paragraphs else None


def set_paragraph_color(paragraph, color: RGBColor) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = color


def append_new_items(doc: Document, items: Sequence[Dict[str, object]]) -> int:
    normalized_items: List[Dict[str, object]] = []
    for row in items:
        if not isinstance(row, dict):
            continue
        item = normalize_new_item(row)
        if item is None:
            continue
        if item["status"] != "filled":
            continue
        normalized_items.append(item)
    if not normalized_items:
        return 0

    anchor = find_insert_anchor_paragraph(doc)
    if anchor is None:
        return 0

    intro = anchor.insert_paragraph_before("（以下为根据转写补充的临时新增访谈问题）")
    set_paragraph_color(intro, RED)

    inserted = 0
    for idx, item in enumerate(normalized_items, start=1):
        q_text = str(item.get("question_text", "")).strip()
        ans = str(item.get("answer_draft", "")).strip()
        if not q_text or not ans:
            continue
        note_required = bool(item.get("note_required", False))
        note_text = str(item.get("note_text", "") or "").strip()
        if note_required:
            if not note_text:
                note_text = LOW_CONF_NOTE
            if note_text not in ans:
                ans = f"{ans}{note_text}"

        qp = anchor.insert_paragraph_before(f"临时问题{idx}：{q_text}")
        set_paragraph_color(qp, RED)

        rp = anchor.insert_paragraph_before("")
        prefix = rp.add_run("回复：")
        prefix.font.color.rgb = BLACK
        body = rp.add_run(ans)
        body.font.color.rgb = RED
        inserted += 1
    return inserted


def apply_results(args: argparse.Namespace) -> None:
    summary_doc_path = Path(args.summary_docx)
    questions_path = Path(args.questions_json)
    results_path = Path(args.results_json)
    output_docx = Path(args.output_docx)

    anchors_raw = json.loads(questions_path.read_text(encoding="utf-8"))
    results_payload = json.loads(results_path.read_text(encoding="utf-8"))
    results_raw = results_payload.get("results", [])
    new_items_raw = results_payload.get("new_items", [])
    if not isinstance(anchors_raw, list):
        raise ValueError("questions_json must be an array.")
    if not isinstance(results_raw, list):
        raise ValueError("results_json must include results array.")
    if new_items_raw is not None and not isinstance(new_items_raw, list):
        raise ValueError("results_json field new_items must be an array if provided.")

    result_map: Dict[str, Dict[str, object]] = {}
    for item in results_raw:
        if isinstance(item, dict):
            qid = str(item.get("question_id", "")).strip()
            if qid:
                result_map[qid] = item

    doc = Document(str(summary_doc_path))

    applied_count = 0
    for item in anchors_raw:
        if not isinstance(item, dict):
            continue
        qid = str(item.get("question_id", "")).strip()
        rp_idx = int(item.get("reply_paragraph_index", -1))
        if not qid or rp_idx < 0 or rp_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[rp_idx]
        prefix, old_tail = split_reply_prefix_tail(para.text)
        raw_result = result_map.get(
            qid,
            {
                "question_id": qid,
                "status": "blank",
                "confidence": 0,
                "answer_draft": "",
                "note_required": False,
                "note_text": "",
            },
        )
        result = normalize_one_result(qid, raw_result)
        new_tail = build_new_reply_text(result)
        rewrite_reply_paragraph(para, prefix, old_tail, new_tail)
        applied_count += 1

    inserted_new_items = append_new_items(doc, new_items_raw or [])

    ensure_dir(output_docx.parent)
    doc.save(str(output_docx))
    print(f"[apply] updated replies: {applied_count}")
    print(f"[apply] inserted temporary new items: {inserted_new_items}")
    print(f"[apply] wrote: {output_docx}")


def run_pipeline(args: argparse.Namespace) -> None:
    output_dir = Path(args.output_dir)
    default_input = str(Path("work") / "llm_input.json")
    default_output = str(Path("work") / "llm_results.json")
    default_questions = str(Path("work") / "questions.json")
    if args.input_json == default_input:
        args.input_json = str(output_dir / "llm_input.json")
    if args.output_json == default_output:
        args.output_json = str(output_dir / "llm_results.json")
    if args.questions_json == default_questions:
        args.questions_json = str(output_dir / "questions.json")

    prepare(args)
    run_llm(args)
    apply_results(args)


def resolve_agent_pack_defaults(args: argparse.Namespace) -> None:
    output_dir = Path(args.output_dir)
    default_input = str(Path("work") / "llm_input.json")
    default_pack_dir = str(Path("work") / "agent_tasks")
    if getattr(args, "input_json", "") == default_input:
        args.input_json = str(output_dir / "llm_input.json")
    if getattr(args, "pack_dir", "") == default_pack_dir:
        args.pack_dir = str(output_dir / "agent_tasks")


def resolve_finalize_defaults(args: argparse.Namespace) -> None:
    output_dir = Path(args.output_dir)
    default_manifest = str(Path("work") / "agent_tasks" / "manifest.json")
    default_results = str(Path("work") / "llm_results.json")
    default_questions = str(Path("work") / "questions.json")
    if getattr(args, "manifest_json", "") == default_manifest:
        args.manifest_json = str(output_dir / "agent_tasks" / "manifest.json")
    if getattr(args, "results_json", "") == default_results:
        args.results_json = str(output_dir / "llm_results.json")
    if getattr(args, "questions_json", "") == default_questions:
        args.questions_json = str(output_dir / "questions.json")


def clean_output_dir(path: Path) -> None:
    cwd = Path.cwd().resolve()
    target = path.resolve()
    if target in {cwd, cwd.parent, Path("/")}:
        raise ValueError(f"Refusing to clean unsafe path: {target}")
    if cwd != target and cwd not in target.parents:
        raise ValueError(
            f"--clean only supports output-dir inside current workspace: {cwd}"
        )
    if path.exists():
        shutil.rmtree(path)
        print(f"[prepare-agent] cleaned existing output-dir: {path}")


def prepare_agent(args: argparse.Namespace) -> None:
    if bool(getattr(args, "clean", False)):
        clean_output_dir(Path(args.output_dir))
    prepare(args)
    resolve_agent_pack_defaults(args)
    agent_pack(args)


def finalize_agent(args: argparse.Namespace) -> None:
    resolve_finalize_defaults(args)
    args.output_json = args.results_json
    merge_agent_results(args)
    apply_results(args)


def parser_with_common() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description=(
            "访谈纪要融合自动化脚本: "
            "prepare / run-llm / apply / pipeline / "
            "agent-pack / merge-agent-results / prepare-agent / finalize-agent"
        )
    )
    sub = p.add_subparsers(dest="command", required=True)

    def add_common_io(sp):
        sp.add_argument(
            "--transcript-docx",
            default="转写.docx",
            help="path to transcript docx",
        )
        sp.add_argument(
            "--summary-docx",
            default="访谈纪要.docx",
            help="path to summary docx",
        )
        sp.add_argument(
            "--output-dir",
            default="work",
            help="workspace folder for temp files",
        )
        sp.add_argument(
            "--project-id",
            default="",
            help="project id for llm payload (default: summary docx stem)",
        )

    sp_prepare = sub.add_parser("prepare", help="convert and extract intermediate files")
    add_common_io(sp_prepare)
    sp_prepare.set_defaults(func=prepare)

    sp_llm = sub.add_parser("run-llm", help="call model API and save results json")
    sp_llm.add_argument("--input-json", default="work/llm_input.json")
    sp_llm.add_argument("--output-json", default="work/llm_results.json")
    sp_llm.add_argument("--api-key", default="", help="API key (fallback: OPENAI_API_KEY)")
    sp_llm.add_argument("--base-url", default="", help="OpenAI-compatible base URL")
    sp_llm.add_argument("--model", default="", help="model name")
    sp_llm.add_argument("--batch-size", type=int, default=25)
    sp_llm.add_argument("--temperature", type=float, default=0.1)
    sp_llm.add_argument("--timeout", type=int, default=120)
    sp_llm.add_argument("--max-retries", type=int, default=3)
    sp_llm.set_defaults(func=run_llm)

    sp_apply = sub.add_parser("apply", help="write results json back into docx")
    sp_apply.add_argument("--summary-docx", default="访谈纪要.docx")
    sp_apply.add_argument("--questions-json", default="work/questions.json")
    sp_apply.add_argument("--results-json", default="work/llm_results.json")
    sp_apply.add_argument(
        "--output-docx",
        default="访谈纪要完成版ClaudeQwen.docx",
        help="output merged docx",
    )
    sp_apply.set_defaults(func=apply_results)

    sp_pipeline = sub.add_parser("pipeline", help="run prepare + run-llm + apply")
    add_common_io(sp_pipeline)
    sp_pipeline.add_argument("--input-json", default="work/llm_input.json")
    sp_pipeline.add_argument("--output-json", default="work/llm_results.json")
    sp_pipeline.add_argument("--questions-json", default="work/questions.json")
    sp_pipeline.add_argument(
        "--output-docx",
        default="访谈纪要完成版ClaudeQwen.docx",
        help="output merged docx",
    )
    sp_pipeline.add_argument("--api-key", default="", help="API key (fallback: OPENAI_API_KEY)")
    sp_pipeline.add_argument("--base-url", default="", help="OpenAI-compatible base URL")
    sp_pipeline.add_argument("--model", default="", help="model name")
    sp_pipeline.add_argument("--batch-size", type=int, default=25)
    sp_pipeline.add_argument("--temperature", type=float, default=0.1)
    sp_pipeline.add_argument("--timeout", type=int, default=120)
    sp_pipeline.add_argument("--max-retries", type=int, default=3)
    sp_pipeline.set_defaults(func=run_pipeline)

    sp_agent_pack = sub.add_parser(
        "agent-pack", help="generate Codex/Claude batch tasks from llm_input.json"
    )
    sp_agent_pack.add_argument("--input-json", default="work/llm_input.json")
    sp_agent_pack.add_argument("--pack-dir", default="work/agent_tasks")
    sp_agent_pack.add_argument("--batch-size", type=int, default=25)
    sp_agent_pack.set_defaults(func=agent_pack)

    sp_merge = sub.add_parser(
        "merge-agent-results",
        help="merge batch output json files into unified llm_results.json",
    )
    sp_merge.add_argument("--manifest-json", default="work/agent_tasks/manifest.json")
    sp_merge.add_argument("--output-json", default="work/llm_results.json")
    sp_merge.add_argument(
        "--strict",
        action="store_true",
        help="fail if any batch output file is missing",
    )
    sp_merge.set_defaults(func=merge_agent_results)

    sp_prepare_agent = sub.add_parser(
        "prepare-agent", help="run prepare + agent-pack (recommended for Codex/Claude)"
    )
    add_common_io(sp_prepare_agent)
    sp_prepare_agent.add_argument("--input-json", default="work/llm_input.json")
    sp_prepare_agent.add_argument("--pack-dir", default="work/agent_tasks")
    sp_prepare_agent.add_argument("--batch-size", type=int, default=25)
    sp_prepare_agent.add_argument(
        "--clean",
        action="store_true",
        help="clean output-dir before generating new artifacts",
    )
    sp_prepare_agent.set_defaults(func=prepare_agent)

    sp_finalize_agent = sub.add_parser(
        "finalize-agent",
        help="run merge-agent-results + apply (after agent produced batch outputs)",
    )
    sp_finalize_agent.add_argument("--output-dir", default="work")
    sp_finalize_agent.add_argument("--summary-docx", default="访谈纪要.docx")
    sp_finalize_agent.add_argument("--manifest-json", default="work/agent_tasks/manifest.json")
    sp_finalize_agent.add_argument("--results-json", default="work/llm_results.json")
    sp_finalize_agent.add_argument("--questions-json", default="work/questions.json")
    sp_finalize_agent.add_argument(
        "--output-docx",
        default="访谈纪要完成版ClaudeQwen.docx",
        help="output merged docx",
    )
    sp_finalize_agent.add_argument(
        "--strict",
        action="store_true",
        help="fail if any batch output file is missing",
    )
    sp_finalize_agent.set_defaults(func=finalize_agent)

    return p


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = parser_with_common()
    args = parser.parse_args(argv)
    try:
        args.func(args)
    except Exception as exc:  # pylint: disable=broad-except
        print(f"[error] {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
